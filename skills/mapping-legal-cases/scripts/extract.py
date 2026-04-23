#!/usr/bin/env python3
"""
mapping-legal-cases: Universal document extractor.

Environment-aware: detects which tools are installed and adapts.
- Always works: plain text, markdown, CSV, JSON
- With pypdf/PyMuPDF: PDF text extraction
- With python-docx/pandoc: Word documents
- With openpyxl: Excel
- With Pillow: image EXIF metadata
- With ffprobe: audio/video metadata
- With tesseract: OCR for scanned PDFs and images (auto-used if available)
- With whisper/faster-whisper: audio transcription (auto-used if available)
- Graceful degradation: unsupported files flagged for manual review

Usage:
    python extract.py <file>                        # extract single file, JSON output
    python extract.py --preflight                   # check installed tools
    python extract.py --scan <folder>               # list files by category
    python extract.py <file> --cache DIR --case-root ROOT   # with caching
    python extract.py <file> --max-chars N          # truncate output

Output: JSON on stdout. Returns non-zero exit on failure.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import subprocess
import sys
from pathlib import Path


# File category mapping
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".heic", ".heif", ".bmp", ".gif", ".webp"}
AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".aac", ".ogg", ".flac", ".wma", ".opus"}
VIDEO_EXTS = {".mp4", ".mov", ".avi", ".mkv", ".wmv", ".webm", ".flv", ".m4v"}
PANDOC_EXTS = {".docx", ".odt", ".rtf", ".html", ".htm", ".epub"}
MARKDOWN_EXTS = {".md", ".markdown"}
PLAIN_TEXT_EXTS = {".txt", ".log"}
CSV_EXTS = {".csv", ".tsv"}
JSON_EXTS = {".json"}
EMAIL_EXTS = {".eml", ".msg"}
SPREADSHEET_EXTS = {".xlsx", ".xls", ".xlsm"}
ARCHIVE_EXTS = {".zip"}


def _result(file_path: str, ext: str, size: int) -> dict:
    """Default result shape."""
    return {
        "success": False,
        "file": file_path,
        "ext": ext,
        "size_bytes": size,
        "tool_used": "",
        "text": "",
        "char_count": 0,
        "page_count": None,
        "language_detected": None,
        "warnings": [],
        "needs_ocr": False,
        "needs_transcription": False,
        "needs_video_analysis": False,
        "metadata": {},
    }


def _detect_language(sample: str) -> str | None:
    """Cheap language detection based on Unicode script counts."""
    if not sample:
        return None
    counts = {"hebrew": 0, "arabic": 0, "cyrillic": 0, "cjk": 0, "latin": 0}
    for ch in sample[:4000]:
        code = ord(ch)
        if 0x0590 <= code <= 0x05FF:
            counts["hebrew"] += 1
        elif 0x0600 <= code <= 0x06FF:
            counts["arabic"] += 1
        elif 0x0400 <= code <= 0x04FF:
            counts["cyrillic"] += 1
        elif 0x4E00 <= code <= 0x9FFF:
            counts["cjk"] += 1
        elif ch.isalpha() and code < 128:
            counts["latin"] += 1
    if sum(counts.values()) == 0:
        return None
    lang = max(counts, key=counts.get)
    mapping = {"hebrew": "he", "arabic": "ar", "cyrillic": "ru", "cjk": "zh", "latin": "en"}
    return mapping[lang]


def _truncate(text: str, max_chars: int | None) -> str:
    if max_chars and len(text) > max_chars:
        return text[:max_chars] + f"\n\n... [truncated at {max_chars:,} chars of {len(text):,}]"
    return text


def _tool_available(name: str) -> bool:
    """Check if a command is available on PATH."""
    try:
        proc = subprocess.run(
            ["where", name] if os.name == "nt" else ["which", name],
            capture_output=True, text=True, timeout=3,
        )
        return proc.returncode == 0
    except Exception:
        return False


def _module_available(module: str) -> bool:
    try:
        __import__(module)
        return True
    except ImportError:
        return False


# ---------- PDF extraction (multiple backend options) ----------

def extract_pdf(path: Path, result: dict) -> dict:
    """Try PyMuPDF first, fall back to pypdf, then to pdftotext CLI."""
    # Try PyMuPDF (best quality)
    if _module_available("fitz"):
        try:
            import fitz
            doc = fitz.open(str(path))
            result["page_count"] = doc.page_count
            parts = [f"[Page {i+1}]\n{page.get_text('text') or ''}" for i, page in enumerate(doc)]
            doc.close()
            full_text = "\n\n".join(p for p in parts if p.strip().splitlines()[1:])
            if full_text.strip():
                result["text"] = full_text
                result["char_count"] = len(full_text)
                result["tool_used"] = "PyMuPDF"
                result["success"] = True
                _maybe_extract_tables(path, result)
                return result
            # No text - likely scanned
            result["warnings"].append("PDF has no extractable text (likely scanned)")
            result["needs_ocr"] = True
            result["tool_used"] = "PyMuPDF (empty)"
            _try_ocr(path, result)
            return result
        except Exception as e:
            result["warnings"].append(f"PyMuPDF failed: {e}")

    # Fallback to pypdf
    if _module_available("pypdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            result["page_count"] = len(reader.pages)
            parts = [f"[Page {i+1}]\n{page.extract_text() or ''}" for i, page in enumerate(reader.pages)]
            full_text = "\n\n".join(parts)
            if full_text.strip():
                result["text"] = full_text
                result["char_count"] = len(full_text)
                result["tool_used"] = "pypdf"
                result["success"] = True
                return result
            result["warnings"].append("PDF has no extractable text (likely scanned)")
            result["needs_ocr"] = True
            result["tool_used"] = "pypdf (empty)"
            _try_ocr(path, result)
            return result
        except Exception as e:
            result["warnings"].append(f"pypdf failed: {e}")

    # Fallback to pdftotext CLI
    if _tool_available("pdftotext"):
        try:
            proc = subprocess.run(
                ["pdftotext", "-layout", str(path), "-"],
                capture_output=True, text=True, encoding="utf-8", timeout=60,
            )
            if proc.returncode == 0 and proc.stdout.strip():
                result["text"] = proc.stdout
                result["char_count"] = len(proc.stdout)
                result["tool_used"] = "pdftotext"
                result["success"] = True
                return result
        except Exception as e:
            result["warnings"].append(f"pdftotext failed: {e}")

    result["warnings"].append(
        "No PDF extraction backend available. Install one: "
        "pip install pymupdf (recommended) OR pip install pypdf OR install poppler-utils"
    )
    result["needs_ocr"] = True
    return result


def _maybe_extract_tables(path: Path, result: dict) -> None:
    """Add pdfplumber tables if the text seems thin and pdfplumber is available."""
    if not _module_available("pdfplumber"):
        return
    if result["char_count"] > 2000:
        return
    try:
        import pdfplumber
        extra_parts = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                for t in page.extract_tables() or []:
                    rows = ["| " + " | ".join(str(c) if c is not None else "" for c in r) + " |" for r in t]
                    extra_parts.append(f"[Page {i+1} Table]\n" + "\n".join(rows))
        if extra_parts:
            result["text"] += "\n\n" + "\n\n".join(extra_parts)
            result["char_count"] = len(result["text"])
            result["tool_used"] += "+pdfplumber"
    except Exception as e:
        result["warnings"].append(f"pdfplumber table extraction skipped: {e}")


def _try_ocr(path: Path, result: dict) -> None:
    """Optional: if tesseract + pdf2image available, OCR the PDF."""
    if not (_tool_available("tesseract") and _module_available("pdf2image")):
        return
    try:
        from pdf2image import convert_from_path
        import pytesseract
        images = convert_from_path(str(path))
        parts = []
        for i, img in enumerate(images):
            text = pytesseract.image_to_string(img)
            if text.strip():
                parts.append(f"[Page {i+1} OCR]\n{text}")
        if parts:
            result["text"] = "\n\n".join(parts)
            result["char_count"] = len(result["text"])
            result["tool_used"] = "tesseract-ocr"
            result["needs_ocr"] = False
            result["success"] = True
    except Exception as e:
        result["warnings"].append(f"OCR attempt failed: {e}")


# ---------- Word documents ----------

def extract_pandoc(path: Path, result: dict) -> dict:
    """Extract via pandoc (DOCX, ODT, RTF, HTML, EPUB)."""
    if _tool_available("pandoc"):
        try:
            proc = subprocess.run(
                ["pandoc", str(path), "-t", "plain", "--wrap=none"],
                capture_output=True, text=True, encoding="utf-8", timeout=60,
            )
            if proc.returncode == 0:
                result["text"] = proc.stdout
                result["char_count"] = len(result["text"])
                result["tool_used"] = "pandoc"
                result["success"] = True
                return result
            result["warnings"].append(f"pandoc failed: {proc.stderr[:200]}")
        except Exception as e:
            result["warnings"].append(f"pandoc exception: {e}")

    # Fallback: python-docx for .docx only
    if path.suffix.lower() == ".docx" and _module_available("docx"):
        try:
            import docx
            d = docx.Document(str(path))
            parts = [p.text for p in d.paragraphs]
            for tbl in d.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            result["text"] = "\n".join(parts)
            result["char_count"] = len(result["text"])
            result["tool_used"] = "python-docx"
            result["success"] = True
            return result
        except Exception as e:
            result["warnings"].append(f"python-docx failed: {e}")

    result["warnings"].append(
        "No backend for this format. Install pandoc (https://pandoc.org) or python-docx (pip install python-docx)"
    )
    return result


def extract_doc_legacy(path: Path, result: dict) -> dict:
    """Legacy .doc files via LibreOffice, if installed."""
    if _tool_available("soffice"):
        try:
            outdir = path.parent / ".tmp_conv"
            outdir.mkdir(parents=True, exist_ok=True)
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx", "--outdir", str(outdir), str(path)],
                capture_output=True, timeout=120,
            )
            converted = outdir / (path.stem + ".docx")
            if converted.exists():
                r = extract_pandoc(converted, result)
                r["tool_used"] = "libreoffice+" + r.get("tool_used", "pandoc")
                try:
                    converted.unlink()
                    outdir.rmdir()
                except OSError:
                    pass
                return r
        except Exception as e:
            result["warnings"].append(f"libreoffice conversion failed: {e}")
    result["warnings"].append(
        "Legacy .doc format needs LibreOffice for auto-conversion. "
        "Install from https://libreoffice.org or manually convert to .docx"
    )
    return result


# ---------- Spreadsheets ----------

def extract_xlsx(path: Path, result: dict) -> dict:
    if not _module_available("openpyxl"):
        result["warnings"].append("openpyxl not installed (pip install openpyxl)")
        return result
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            parts.append(f"## Sheet: {sheet}\n")
            for row in ws.iter_rows(values_only=True):
                vals = [str(c) if c is not None else "" for c in row]
                if any(v.strip() for v in vals):
                    parts.append(" | ".join(vals))
            parts.append("")
        wb.close()
        result["text"] = "\n".join(parts)
        result["char_count"] = len(result["text"])
        result["tool_used"] = "openpyxl"
        result["success"] = True
    except Exception as e:
        result["warnings"].append(f"xlsx extraction failed: {e}")
    return result


# ---------- Email ----------

def extract_email(path: Path, result: dict) -> dict:
    import email
    from email import policy

    if path.suffix.lower() == ".msg":
        if _module_available("extract_msg"):
            try:
                import extract_msg
                msg = extract_msg.Message(str(path))
                header = f"From: {msg.sender}\nTo: {msg.to}\nDate: {msg.date}\nSubject: {msg.subject}\n\n"
                result["text"] = header + (msg.body or "")
                result["char_count"] = len(result["text"])
                result["tool_used"] = "extract_msg"
                result["success"] = True
                return result
            except Exception as e:
                result["warnings"].append(f"extract_msg failed: {e}")
        result["warnings"].append("Install extract-msg for full .msg support (pip install extract-msg)")
        return result

    # .eml
    try:
        with open(path, "rb") as f:
            msg = email.message_from_binary_file(f, policy=policy.default)
        header = (
            f"From: {msg.get('From', '')}\n"
            f"To: {msg.get('To', '')}\n"
            f"Cc: {msg.get('Cc', '')}\n"
            f"Date: {msg.get('Date', '')}\n"
            f"Subject: {msg.get('Subject', '')}\n\n"
        )
        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                ct = part.get_content_type()
                if ct == "text/plain":
                    try:
                        body += part.get_content() + "\n"
                    except Exception:
                        pass
                elif ct == "text/html" and not body:
                    try:
                        body += part.get_content()
                    except Exception:
                        pass
        else:
            body = msg.get_content()
        result["text"] = header + body
        result["char_count"] = len(result["text"])
        result["tool_used"] = "email-stdlib"
        result["success"] = True
    except Exception as e:
        result["warnings"].append(f"email extraction failed: {e}")
    return result


# ---------- Plain text formats ----------

def extract_plain_text(path: Path, result: dict) -> dict:
    """Try multiple encodings for best legacy-file compatibility."""
    encodings = ["utf-8", "utf-8-sig", "cp1255", "cp1256", "cp1252", "windows-1251", "latin-1"]
    for enc in encodings:
        try:
            result["text"] = path.read_text(encoding=enc)
            result["char_count"] = len(result["text"])
            result["tool_used"] = f"text-read ({enc})"
            result["success"] = True
            return result
        except UnicodeDecodeError:
            continue
    result["warnings"].append("Could not decode text with any common encoding")
    return result


def extract_json(path: Path, result: dict) -> dict:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        result["text"] = json.dumps(data, indent=2, ensure_ascii=False)
        result["char_count"] = len(result["text"])
        result["tool_used"] = "json-parse"
        result["success"] = True
        result["metadata"] = {"type": "json", "root_type": type(data).__name__}
    except Exception as e:
        result["warnings"].append(f"json parse failed: {e}")
        # Fallback to raw read
        return extract_plain_text(path, result)
    return result


# ---------- Media (metadata only, optional OCR/transcription) ----------

def extract_image(path: Path, result: dict) -> dict:
    result["tool_used"] = "metadata-only"
    result["needs_ocr"] = True
    meta = {"note": "Image file - OCR required for text content extraction"}

    if _module_available("PIL"):
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS
            img = Image.open(path)
            meta["size"] = f"{img.width}x{img.height}"
            meta["mode"] = img.mode
            try:
                exif = img._getexif() or {}
                for tag_id, value in exif.items():
                    tag = TAGS.get(tag_id, str(tag_id))
                    if tag in ("DateTime", "DateTimeOriginal", "DateTimeDigitized"):
                        meta[tag] = str(value)
            except Exception:
                pass
            img.close()
        except Exception as e:
            result["warnings"].append(f"image metadata: {e}")

    # If tesseract available, try OCR
    if _tool_available("tesseract"):
        try:
            proc = subprocess.run(
                ["tesseract", str(path), "-", "-l", "eng+heb+ara+rus"],
                capture_output=True, text=True, encoding="utf-8", timeout=60,
            )
            if proc.returncode == 0 and proc.stdout.strip():
                result["text"] = f"[Image OCR: {path.name}]\n{proc.stdout}"
                result["char_count"] = len(result["text"])
                result["tool_used"] = "tesseract-ocr"
                result["needs_ocr"] = False
                result["metadata"] = meta
                result["success"] = True
                return result
        except Exception as e:
            result["warnings"].append(f"tesseract attempt failed: {e}")

    result["metadata"] = meta
    summary = f"[Image: {path.name}]\nDimensions: {meta.get('size', 'unknown')}\nEXIF date: {meta.get('DateTimeOriginal', meta.get('DateTime', 'none'))}\n\nNEEDS OCR for text content."
    result["text"] = summary
    result["char_count"] = len(summary)
    result["success"] = True
    return result


def extract_audio(path: Path, result: dict) -> dict:
    result["tool_used"] = "metadata-only"
    result["needs_transcription"] = True
    meta = _ffprobe_metadata(path)
    result["metadata"] = meta

    # Try whisper or faster-whisper if available
    transcription = _try_transcription(path, result)
    if transcription:
        result["text"] = f"[Audio transcription: {path.name}]\n{transcription}"
        result["char_count"] = len(result["text"])
        result["needs_transcription"] = False
        result["success"] = True
        return result

    summary = f"[Audio: {path.name}]\nDuration: {meta.get('duration', 'unknown')}\nFormat: {meta.get('format', 'unknown')}\n\nNEEDS TRANSCRIPTION for content."
    result["text"] = summary
    result["char_count"] = len(summary)
    result["success"] = True
    return result


def extract_video(path: Path, result: dict) -> dict:
    result["tool_used"] = "metadata-only"
    result["needs_video_analysis"] = True
    result["needs_transcription"] = True
    meta = _ffprobe_metadata(path)
    result["metadata"] = meta

    # Audio transcription may still work on video files
    transcription = _try_transcription(path, result)
    if transcription:
        result["text"] = f"[Video with audio transcription: {path.name}]\n{transcription}\n\n(Visual analysis still pending.)"
        result["char_count"] = len(result["text"])
        result["needs_transcription"] = False
        result["success"] = True
        return result

    summary = f"[Video: {path.name}]\nDuration: {meta.get('duration', 'unknown')}\nFormat: {meta.get('format', 'unknown')}\n\nNEEDS VIDEO ANALYSIS + TRANSCRIPTION."
    result["text"] = summary
    result["char_count"] = len(summary)
    result["success"] = True
    return result


def _try_transcription(path: Path, result: dict) -> str | None:
    """Attempt transcription with faster-whisper or whisper if installed."""
    if _module_available("faster_whisper"):
        try:
            from faster_whisper import WhisperModel
            model = WhisperModel("base", device="cpu", compute_type="int8")
            segments, _info = model.transcribe(str(path))
            text = " ".join(seg.text for seg in segments)
            if text.strip():
                result["tool_used"] = "faster-whisper"
                return text
        except Exception as e:
            result["warnings"].append(f"faster-whisper failed: {e}")

    if _module_available("whisper"):
        try:
            import whisper
            model = whisper.load_model("base")
            out = model.transcribe(str(path))
            text = out.get("text", "")
            if text.strip():
                result["tool_used"] = "openai-whisper"
                return text
        except Exception as e:
            result["warnings"].append(f"whisper failed: {e}")
    return None


def _ffprobe_metadata(path: Path) -> dict:
    if not _tool_available("ffprobe"):
        return {"note": "ffprobe not available - install ffmpeg for media metadata"}
    try:
        proc = subprocess.run(
            ["ffprobe", "-v", "error", "-show_format", "-show_streams", "-of", "json", str(path)],
            capture_output=True, text=True, timeout=15,
        )
        if proc.returncode == 0:
            data = json.loads(proc.stdout)
            fmt = data.get("format", {})
            return {
                "duration": fmt.get("duration"),
                "format": fmt.get("format_name"),
                "size": fmt.get("size"),
                "streams": len(data.get("streams", [])),
            }
    except Exception:
        pass
    return {"note": "ffprobe failed"}


def extract_archive(path: Path, result: dict) -> dict:
    try:
        import zipfile
        with zipfile.ZipFile(path) as z:
            names = z.namelist()
        parts = [f"[Archive: {path.name}]", f"Contains {len(names)} entries:"]
        parts += [f"- {n}" for n in names[:100]]
        if len(names) > 100:
            parts.append(f"... and {len(names) - 100} more")
        result["text"] = "\n".join(parts)
        result["char_count"] = len(result["text"])
        result["tool_used"] = "zipfile-list"
        result["success"] = True
        result["warnings"].append("Archive contents not extracted - unzip manually if needed")
    except Exception as e:
        result["warnings"].append(f"archive listing failed: {e}")
    return result


# ---------- Router ----------

def extract(path: Path, max_chars: int | None = None) -> dict:
    size = path.stat().st_size if path.exists() else 0
    ext = path.suffix.lower()
    result = _result(str(path), ext, size)

    if not path.exists():
        result["warnings"].append("File does not exist")
        return result

    if size == 0:
        result["warnings"].append("Empty file")
        result["tool_used"] = "none"
        result["success"] = True
        return result

    if ext == ".pdf":
        extract_pdf(path, result)
    elif ext == ".doc":
        extract_doc_legacy(path, result)
    elif ext in PANDOC_EXTS:
        extract_pandoc(path, result)
    elif ext in SPREADSHEET_EXTS:
        extract_xlsx(path, result)
    elif ext in EMAIL_EXTS:
        extract_email(path, result)
    elif ext in MARKDOWN_EXTS:
        extract_plain_text(path, result)
    elif ext in JSON_EXTS:
        extract_json(path, result)
    elif ext in PLAIN_TEXT_EXTS or ext in CSV_EXTS:
        extract_plain_text(path, result)
    elif ext in IMAGE_EXTS:
        extract_image(path, result)
    elif ext in AUDIO_EXTS:
        extract_audio(path, result)
    elif ext in VIDEO_EXTS:
        extract_video(path, result)
    elif ext in ARCHIVE_EXTS:
        extract_archive(path, result)
    else:
        # Try plain-text decode for unknown extensions (many legal files have no extension)
        tried = extract_plain_text(path, result)
        if not tried["success"]:
            result["warnings"].append(f"Unsupported extension: {ext}")
            result["tool_used"] = "unsupported"

    if result["text"]:
        result["language_detected"] = _detect_language(result["text"])
        result["text"] = _truncate(result["text"], max_chars)

    return result


# ---------- Preflight ----------

def preflight() -> dict:
    """Report tool/module availability and give install hints."""
    status = {
        "python": sys.version.split()[0],
        "platform": sys.platform,
        "modules": {
            "pymupdf": _module_available("fitz"),
            "pypdf": _module_available("pypdf"),
            "pdfplumber": _module_available("pdfplumber"),
            "python_docx": _module_available("docx"),
            "openpyxl": _module_available("openpyxl"),
            "pillow": _module_available("PIL"),
            "extract_msg": _module_available("extract_msg"),
            "pdf2image": _module_available("pdf2image"),
            "pytesseract": _module_available("pytesseract"),
            "whisper": _module_available("whisper"),
            "faster_whisper": _module_available("faster_whisper"),
        },
        "tools": {
            "pandoc": _tool_available("pandoc"),
            "pdftotext": _tool_available("pdftotext"),
            "ffprobe": _tool_available("ffprobe"),
            "ffmpeg": _tool_available("ffmpeg"),
            "tesseract": _tool_available("tesseract"),
            "libreoffice": _tool_available("soffice"),
        },
    }

    # Capabilities summary - what can we actually do?
    caps = status["capabilities"] = {
        "pdf_text": status["modules"]["pymupdf"] or status["modules"]["pypdf"] or status["tools"]["pdftotext"],
        "word": status["tools"]["pandoc"] or status["modules"]["python_docx"],
        "word_legacy": status["tools"]["libreoffice"] and status["tools"]["pandoc"],
        "excel": status["modules"]["openpyxl"],
        "email_eml": True,  # stdlib
        "email_msg": status["modules"]["extract_msg"],
        "image_metadata": status["modules"]["pillow"],
        "ocr_images": status["tools"]["tesseract"],
        "ocr_scanned_pdfs": status["tools"]["tesseract"] and status["modules"]["pdf2image"],
        "audio_metadata": status["tools"]["ffprobe"],
        "video_metadata": status["tools"]["ffprobe"],
        "transcription": status["modules"]["whisper"] or status["modules"]["faster_whisper"],
        "plain_text": True,
        "csv": True,
        "json": True,
        "markdown": True,
    }

    # A "ready" status requires minimum: can handle PDFs or Word OR we only need text/CSV/JSON
    status["ready"] = caps["plain_text"]  # we can always at least do text
    status["fully_ready"] = caps["pdf_text"] and caps["word"] and caps["excel"]

    hints = []
    if not caps["pdf_text"]:
        hints.append("Install PDF support: pip install pymupdf  (recommended) OR pip install pypdf")
    if not caps["word"]:
        hints.append("Install Word support: install pandoc (https://pandoc.org) OR pip install python-docx")
    if not caps["excel"]:
        hints.append("Install Excel support: pip install openpyxl")
    if not caps["image_metadata"]:
        hints.append("Optional - image metadata: pip install Pillow")
    if not caps["email_msg"]:
        hints.append("Optional - Outlook .msg: pip install extract-msg")
    if not caps["ocr_images"]:
        hints.append("Optional - OCR for images/scanned PDFs: install tesseract-ocr")
    if not caps["transcription"]:
        hints.append("Optional - audio/video transcription: pip install faster-whisper  (or openai-whisper)")
    if not caps["audio_metadata"]:
        hints.append("Optional - audio/video metadata: install ffmpeg")

    status["install_hints"] = hints
    return status


# ---------- Folder scan ----------

def scan_folder(folder: Path) -> dict:
    EXCLUDE_DIRS = {".casebase", ".git", "node_modules", "__pycache__", ".venv", "venv", ".svn", ".hg"}
    files = []
    for p in folder.rglob("*"):
        if p.is_dir():
            continue
        if any(part in EXCLUDE_DIRS for part in p.parts):
            continue
        if p.name.startswith("~$"):  # Office lock files
            continue
        ext = p.suffix.lower()
        category = "other"
        if ext == ".pdf":
            category = "pdf"
        elif ext in PANDOC_EXTS or ext == ".doc":
            category = "word"
        elif ext in MARKDOWN_EXTS:
            category = "markdown"
        elif ext in SPREADSHEET_EXTS:
            category = "spreadsheet"
        elif ext in EMAIL_EXTS:
            category = "email"
        elif ext in IMAGE_EXTS:
            category = "image"
        elif ext in AUDIO_EXTS:
            category = "audio"
        elif ext in VIDEO_EXTS:
            category = "video"
        elif ext in PLAIN_TEXT_EXTS:
            category = "text"
        elif ext in CSV_EXTS:
            category = "csv"
        elif ext in JSON_EXTS:
            category = "json"
        elif ext in ARCHIVE_EXTS:
            category = "archive"
        try:
            size = p.stat().st_size
            mtime = p.stat().st_mtime
        except Exception:
            size = 0
            mtime = 0
        files.append({
            "path": str(p.relative_to(folder)).replace("\\", "/"),
            "ext": ext,
            "category": category,
            "size": size,
            "mtime": mtime,
        })

    counts, total_size = {}, 0
    for f in files:
        counts[f["category"]] = counts.get(f["category"], 0) + 1
        total_size += f["size"]

    return {
        "total_files": len(files),
        "total_size_bytes": total_size,
        "by_category": counts,
        "files": files,
    }


# ---------- Cache ----------

def cache_path_for(file_path: Path, cache_dir: Path, case_root: Path) -> Path:
    rel = str(file_path.resolve().relative_to(case_root.resolve())).replace("\\", "/")
    h = hashlib.sha1(rel.encode("utf-8")).hexdigest()[:16]
    return cache_dir / f"{h}.json"


# ---------- CLI ----------

def main():
    parser = argparse.ArgumentParser(
        description="mapping-legal-cases: Universal document extractor for AI agents",
    )
    parser.add_argument("file", nargs="?", help="Path to file to extract")
    parser.add_argument("--max-chars", type=int, default=None, help="Truncate text to N chars")
    parser.add_argument("--preflight", action="store_true", help="Check environment tools")
    parser.add_argument("--scan", metavar="FOLDER", help="Scan folder recursively")
    parser.add_argument("--cache", metavar="DIR", help="Write result to cache")
    parser.add_argument("--case-root", metavar="DIR", help="Case root for cache hashing")
    parser.add_argument("--pretty", action="store_true", help="Pretty JSON output")
    args = parser.parse_args()

    if args.preflight:
        out = preflight()
        print(json.dumps(out, indent=2 if args.pretty else None, ensure_ascii=False))
        sys.exit(0 if out["ready"] else 1)

    if args.scan:
        folder = Path(args.scan)
        if not folder.is_dir():
            print(json.dumps({"error": f"Not a directory: {folder}"}), file=sys.stderr)
            sys.exit(1)
        out = scan_folder(folder)
        print(json.dumps(out, indent=2 if args.pretty else None, ensure_ascii=False))
        sys.exit(0)

    if not args.file:
        parser.print_help()
        sys.exit(2)

    path = Path(args.file)
    result = extract(path, max_chars=args.max_chars)

    if args.cache and args.case_root:
        cache_dir = Path(args.cache)
        cache_dir.mkdir(parents=True, exist_ok=True)
        cp = cache_path_for(path, cache_dir, Path(args.case_root))
        with open(cp, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        result["_cached_at"] = str(cp).replace("\\", "/")

    print(json.dumps(result, indent=2 if args.pretty else None, ensure_ascii=False))
    sys.exit(0 if result["success"] else 1)


if __name__ == "__main__":
    main()
